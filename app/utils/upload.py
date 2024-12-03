from fastapi import APIRouter, UploadFile, File
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import pytesseract
import nltk
from nltk.corpus import words

router = APIRouter()

nltk.download('words')
english_words = set(words.words())


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_content = await file.read()

    file_extension = file.filename.split('.')[-1].lower()

    if file_extension == 'pdf':
        text = extract_text_from_pdf(file_content)
    elif file_extension == 'docx':
        text = extract_text_from_docx(file_content)
    elif file_extension in ['jpg', 'jpeg', 'png']:
        text = extract_text_from_image(file_content)
    else:
        return {"error": "Unsupported file type"}

    sanitized_text = sanitize_text(text)
    return {"text": sanitized_text}


def extract_text_from_pdf(file_content):
    doc = fitz.Document(stream=file_content, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text


def extract_text_from_docx(file_content):
    doc = Document(BytesIO(file_content))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def extract_text_from_image(file_content):
    image = Image.open(BytesIO(file_content))
    text = pytesseract.image_to_string(image)
    return text


def sanitize_text(text):
    words_list = text.split()
    sanitized_words = [
        word for word in words_list if word.lower() in english_words]
    return ' '.join(sanitized_words)
